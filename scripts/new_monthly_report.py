"""
Aarise SEO — Monthly Report Generator (v2)

Rebuilt report template addressing the specific gaps identified in the June 2026 report review:
  - One combined report for both sites, not two separate files
  - Real GSC figures pulled directly from the xlsx exports (no ~approximations)
  - Leads-first ordering; sections that need data this pipeline doesn't have yet
    are marked "DATA PENDING" rather than filled with placeholder numbers
  - Country Expansion Tracker: publish status for all 18 target countries x 2 sites
  - Promise Tracker: carries forward last report's recommendations with a status
  - "Waiting on You" is a separate section from the agency's own recommendations

Data sources (repo-relative):
  gsc-data/aarisepharma-MoM-2026-07-23.xlsx       (Queries/Pages/Countries/Devices, current vs prior period)
  gsc-data/aarisehealthcare-MoM-2026-07-23.xlsx   (same, healthcare)
  gsc-data/GA4-snapshot-aarise-pharma.xlsx        (GA4 summary/top pages/channels/cities)
  gsc-data/GA4-snapshot-aarise-health.xlsx        (same, healthcare)

Requires: python-docx (available on /usr/bin/python3 on this machine, not the homebrew ones).
No pandas/openpyxl dependency — xlsx_reader.py is a stdlib-only reader.
"""
import os
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from xlsx_reader import get_sheet_rows  # noqa: E402

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
GSC = os.path.join(REPO, "gsc-data")

PHARMA_MOM = os.path.join(GSC, "aarisepharma-MoM-2026-07-23.xlsx")
HEALTH_MOM = os.path.join(GSC, "aarisehealthcare-MoM-2026-07-23.xlsx")
PHARMA_GA4 = os.path.join(GSC, "GA4-snapshot-aarise-pharma.xlsx")
HEALTH_GA4 = os.path.join(GSC, "GA4-snapshot-aarise-health.xlsx")

CURRENT_LABEL = "20 Jun – 23 Jul 2026"
PREVIOUS_LABEL = "20 May – 20 Jun 2026"

# ── colours ──────────────────────────────────────────────────────────
NAVY = RGBColor(0x0D, 0x1B, 0x2A)
GRN = RGBColor(0x0A, 0x7A, 0x70)
BLU = RGBColor(0x1A, 0x5F, 0xA8)
GREY = RGBColor(0x6B, 0x7A, 0x8D)
DARK = RGBColor(0x3D, 0x4A, 0x5C)
UP = RGBColor(0x1A, 0x7A, 0x4A)
DN = RGBColor(0xB8, 0x30, 0x30)
AMBER = RGBColor(0xB8, 0x86, 0x0A)
HEX = {id(NAVY): "0D1B2A", id(GRN): "0A7A70", id(BLU): "1A5FA8", id(GREY): "6B7A8D",
       id(DARK): "3D4A5C", id(UP): "1A7A4A", id(DN): "B83030", id(AMBER): "B8860A"}


def rhex(c):
    return HEX.get(id(c), "6B7A8D")


# ── docx style helpers (reused from scripts/split_reports.py) ────────
def p(doc, text="", size=10, bold=False, color=None, before=2, after=4, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    if align:
        para.alignment = align
    r = para.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color or DARK
    return para


def mixed(doc, parts, size=10, before=2, after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    for text, bold, color in parts:
        r = para.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color or DARK
    return para


def heading(doc, text, size=18, color=None, before=14, after=4):
    return p(doc, text, size=size, bold=True, color=color or NAVY, before=before, after=after)


def rule(doc, color="CCCCCC"):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "6")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    return para


def section(doc, text, color=None):
    rule(doc, rhex(color or GREY))
    return p(doc, text.upper(), size=8, bold=True, color=color or GREY, before=2, after=6)


def shd(cell, hex6):
    tcPr = cell._tc.get_or_add_tcPr()
    e = OxmlElement("w:shd")
    e.set(qn("w:val"), "clear")
    e.set(qn("w:color"), "auto")
    e.set(qn("w:fill"), hex6)
    tcPr.append(e)


def table(doc, headers, rows, widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        shd(c, "E8ECF2")
        para = c.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        r = para.add_run(h)
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = GREY
        if i > 0:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for ri, row in enumerate(rows):
        bg = "FFFFFF" if ri % 2 == 0 else "F7F8FA"
        for ci, cell in enumerate(row):
            txt = str(cell[0]) if isinstance(cell, tuple) else str(cell)
            bold = cell[1] if isinstance(cell, tuple) and len(cell) > 1 else False
            col = cell[2] if isinstance(cell, tuple) and len(cell) > 2 else None
            c = tbl.rows[ri + 1].cells[ci]
            shd(c, bg)
            para = c.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            r = para.add_run(txt)
            r.font.size = Pt(9)
            r.font.bold = bool(bold)
            r.font.color.rgb = col if isinstance(col, RGBColor) else DARK
            if ci > 0:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if widths:
        for i, w in enumerate(widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def bullet(doc, text, lead=""):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(2)
    if lead:
        r = para.add_run(lead)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
    para.add_run(text).font.size = Pt(10)
    para.runs[-1].font.color.rgb = DARK


def numbered(doc, text, lead=""):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(2)
    if lead:
        r = para.add_run(lead)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
    para.add_run(text).font.size = Pt(10)
    para.runs[-1].font.color.rgb = DARK


def callout(doc, bold_text, body, color=GRN):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.left_indent = Inches(0.15)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    lft = OxmlElement("w:left")
    lft.set(qn("w:val"), "single")
    lft.set(qn("w:sz"), "18")
    lft.set(qn("w:color"), rhex(color))
    pBdr.append(lft)
    pPr.append(pBdr)
    r1 = para.add_run(bold_text + " ")
    r1.font.bold = True
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = color
    para.add_run(body).font.size = Pt(9.5)
    para.runs[-1].font.color.rgb = DARK


def pending_note(doc, text):
    callout(doc, "DATA PENDING —", text, AMBER)


def new_doc():
    doc = Document()
    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(1.0)
    s.top_margin = s.bottom_margin = Inches(1.0)
    return doc


# ── data loading ───────────────────────────────────────────────────────
def load_gsc(path):
    """Returns dict with queries/pages/countries/devices as list-of-dict rows."""
    def rows_to_dicts(rows):
        header = rows[0]
        out = []
        for r in rows[1:]:
            d = {}
            for i, h in enumerate(header):
                d[h] = r[i] if i < len(r) else None
            out.append(d)
        return out

    return {
        "queries": rows_to_dicts(get_sheet_rows(path, "Queries")),
        "pages": rows_to_dicts(get_sheet_rows(path, "Pages")),
        "countries": rows_to_dicts(get_sheet_rows(path, "Countries")),
        "devices": rows_to_dicts(get_sheet_rows(path, "Devices")),
    }


def col(d, prefix, metric):
    """Find a column value by (period-label prefix, metric) since headers embed date ranges."""
    for k, v in d.items():
        if k.startswith(prefix) and k.endswith(metric):
            return v
    return None


def num(v, dec=0):
    if v is None:
        return "0"
    f = float(v)
    if dec == 0:
        return f"{int(round(f)):,}"
    return f"{f:.{dec}f}"


def pct(v):
    if v is None:
        return "0.00%"
    return f"{float(v) * 100:.2f}%"


def load_ga4_summary(path):
    rows = get_sheet_rows(path, "Reports snapshot")
    for i, r in enumerate(rows):
        if r and r[0] == "Active users":
            vals = rows[i + 1]
            return {
                "active_users": vals[0], "new_users": vals[1],
                "avg_engagement": vals[2], "events": vals[3],
            }
    return {}


def load_ga4_top_pages(path, n=8):
    rows = get_sheet_rows(path, "Reports snapshot")
    out = []
    for i, r in enumerate(rows):
        if r and r[0] == "Page title and screen class":
            for row in rows[i + 1:]:
                if not row or not row[0] or row[0].startswith("#"):
                    break
                out.append(row)
                if len(out) >= n:
                    break
            break
    return out


def load_ga4_session_source(path, n=8):
    rows = get_sheet_rows(path, "Reports snapshot")
    out = []
    for i, r in enumerate(rows):
        if r and r[0] == "Session source/medium":
            for row in rows[i + 1:]:
                if not row or not row[0] or row[0].startswith("#"):
                    break
                out.append(row)
                if len(out) >= n:
                    break
            break
    return out


def load_ga4_cities(path, n=8):
    rows = get_sheet_rows(path, "Reports snapshot")
    header_idx = None
    for i, r in enumerate(rows):
        if r and r[0] and "city" in str(r[0]).lower():
            header_idx = i
            break
    if header_idx is None:
        return []
    out = []
    for row in rows[header_idx + 1:]:
        if not row or not row[0] or row[0].startswith("#"):
            break
        out.append(row)
        if len(out) >= n:
            break
    return out


# ── report content builders ───────────────────────────────────────────
def build_leads_section(doc):
    heading(doc, "1. Leads & Enquiries", size=14, before=8)
    pending_note(
        doc,
        "GA4 Key Events (form_start, form_submit, enquiry submissions) require a live GA4 Data API "
        "pull — the static exports in gsc-data/ don't include this breakdown. This section will "
        "populate automatically once the GA4 service account credentials are wired into an automated "
        "pull (the service account already has access to both properties — see repo README/plan for "
        "details). Until then: do not omit this section silently, and do not estimate a number here — "
        "leave it marked pending, as done above, until a real pull is run.",
    )
    p(doc, "What this section will show once connected: total enquiries/form submissions this period, "
           "period-over-period change, and enquiry source (organic/direct/AI-assistant/referral).", size=9, color=GREY)


def build_search_performance(doc, site_label, gsc_data, page_url_col="Top pages"):
    heading(doc, f"2. Search Performance — {site_label}", size=14, before=8)

    section(doc, "Overview", GRN)
    countries = gsc_data["countries"]
    total_clicks_now = sum(float(col(c, CURRENT_LABEL[:5], "Clicks") or 0) for c in countries) if False else None
    # Simplest reliable total: sum device rows (device breakdown always sums to site total)
    devices = gsc_data["devices"]
    def dsum(metric_key_end):
        return sum(float(col(d, "20/06", metric_key_end) or 0) for d in devices)
    clicks_now = dsum("Clicks")
    impr_now = dsum("Impressions")
    devices_prev_clicks = sum(float(col(d, "20/05", "Clicks") or 0) for d in devices)
    devices_prev_impr = sum(float(col(d, "20/05", "Impressions") or 0) for d in devices)
    ctr_now = (clicks_now / impr_now) if impr_now else 0
    ctr_prev = (devices_prev_clicks / devices_prev_impr) if devices_prev_impr else 0

    table(doc, ["Metric", f"This Period ({CURRENT_LABEL})", f"Previous ({PREVIOUS_LABEL})"], [
        [("Clicks", True, NAVY), (num(clicks_now), True, NAVY), (num(devices_prev_clicks), False, GREY)],
        [("Impressions", True, NAVY), (num(impr_now), True, NAVY), (num(devices_prev_impr), False, GREY)],
        [("CTR", False, None), (f"{ctr_now*100:.2f}%", False, None), (f"{ctr_prev*100:.2f}%", False, GREY)],
    ], [2.6, 2.2, 2.2])

    section(doc, "Device Breakdown", GRN)
    rows = []
    for d in devices:
        rows.append([
            (d["Device"], False, None),
            (num(col(d, "20/06", "Clicks")), False, None),
            (num(col(d, "20/06", "Impressions")), False, None),
            (pct(col(d, "20/06", "CTR")), False, None),
            (num(col(d, "20/06", "Position"), 1), False, None),
        ])
    table(doc, ["Device", "Clicks", "Impressions", "CTR", "Avg. Position"], rows, [1.6, 1.3, 1.5, 1.1, 1.5])

    section(doc, "Top Countries", GRN)
    rows = []
    for c in countries[:10]:
        rows.append([
            (c["Country"], False, None),
            (num(col(c, "20/06", "Clicks")), False, None),
            (num(col(c, "20/06", "Impressions")), False, None),
            (pct(col(c, "20/06", "CTR")), False, None),
            (num(col(c, "20/06", "Position"), 1), False, None),
        ])
    table(doc, ["Country", "Clicks", "Impressions", "CTR", "Avg. Position"], rows, [2.0, 1.1, 1.4, 1.0, 1.5])

    section(doc, "Top Queries", GRN)
    rows = []
    for q in gsc_data["queries"][:10]:
        key = list(q.keys())[0]
        rows.append([
            (q[key], False, None),
            (num(col(q, "20/06", "Clicks")), False, None),
            (num(col(q, "20/06", "Impressions")), False, None),
            (pct(col(q, "20/06", "CTR")), False, None),
            (num(col(q, "20/06", "Position"), 1), False, None),
        ])
    table(doc, ["Query", "Clicks", "Impressions", "CTR", "Position"], rows, [2.6, 0.9, 1.2, 0.9, 1.0])

    section(doc, "Top Pages", GRN)
    rows = []
    for pg in gsc_data["pages"][:8]:
        key = list(pg.keys())[0]
        url = pg[key].replace(f"https://{site_label.lower().replace(' ', '')}", "") if pg[key] else ""
        rows.append([
            (pg[key], False, None),
            (num(col(pg, "20/06", "Clicks")), False, None),
            (num(col(pg, "20/06", "Impressions")), False, None),
            (pct(col(pg, "20/06", "CTR")), False, None),
            (num(col(pg, "20/06", "Position"), 1), False, None),
        ])
    table(doc, ["Page", "Clicks", "Impressions", "CTR", "Position"], rows, [3.0, 0.8, 1.1, 0.8, 0.9])


def build_country_tracker(doc, pharma_status, health_status):
    heading(doc, "3. Country Expansion Tracker", size=14, before=8)
    p(doc, "Status of the 18-country expansion initiative, tracked separately from overall site "
           "performance above. Newly published pages show no GSC data yet — impressions typically "
           "take several days to a few weeks to appear after Google crawls a new URL.", size=9.5, color=GREY)

    section(doc, "aarisepharma.com", GRN)
    rows = [[(c, False, None), (s, True, UP if s == "Live — full depth (tier 1-3)" else NAVY)] for c, s in pharma_status]
    table(doc, ["Country", "Status"], rows, [2.5, 4.0])

    section(doc, "aarisehealthcare.com", BLU)
    rows = [[(c, False, None), (s, True, UP if "Live" in s else NAVY)] for c, s in health_status]
    table(doc, ["Country", "Status"], rows, [2.5, 4.0])


def build_ga4_detail(doc, site_label, ga4_summary, top_pages, session_source, cities):
    heading(doc, f"4. GA4 Detail — {site_label}", size=14, before=8)

    section(doc, "Overview", GRN)
    table(doc, ["Metric", "Value"], [
        [("Active Users", True, NAVY), (num(ga4_summary.get("active_users")), True, NAVY)],
        [("New Users", False, None), (num(ga4_summary.get("new_users")), False, None)],
        [("Avg. Engagement / User", False, None),
         (f"{float(ga4_summary.get('avg_engagement', 0)):.1f}s", False, None)],
        [("Event Count", False, None), (num(ga4_summary.get("events")), False, None)],
    ], [3.5, 3.0])

    section(doc, "Traffic by Source/Medium", GRN)
    rows = []
    ai_flag = False
    for s in session_source[:8]:
        label = s[0]
        val = s[1]
        is_ai = "ai-assistant" in label.lower() or "chatgpt" in label.lower() or "gemini" in label.lower() or "perplexity" in label.lower() or "claude" in label.lower()
        if is_ai:
            ai_flag = True
        rows.append([(label, is_ai, GRN if is_ai else None), (num(val), is_ai, GRN if is_ai else None)])
    table(doc, ["Source / Medium", "Sessions"], rows, [4.5, 2.0])
    if ai_flag:
        callout(doc, "AI Traffic:", "AI-assistant sourced sessions are present this period (highlighted above) — "
                                      "worth tracking month over month as a distinct, early-stage channel.", GRN)

    section(doc, "Top Pages (GA4)", GRN)
    rows = [[(r[0], False, None), (num(r[1]), False, None), (num(r[2]), False, None)] for r in top_pages[:8]]
    table(doc, ["Page", "Views", "Active Users"], rows, [4.0, 1.2, 1.5])

    section(doc, "Top Cities", GRN)
    rows = [[(r[0], False, None), (num(r[1]), False, None)] for r in cities[:8]]
    table(doc, ["City", "Active Users"], rows, [4.5, 2.0])


def build_competitor_snapshot(doc):
    heading(doc, "5. Competitor Snapshot", size=14, before=8)
    p(doc, "Directory/listing gap identified this engagement:", size=10)
    bullet(doc, 'Searched "dispersible tablet API manufacturer India" on Pharmacompass — only 8 '
                "results returned, Aarise not among them. Free listing, fastest available lead channel. "
                "Still pending on Sabhya's side.", "Pharmacompass gap: ")
    pending_note(
        doc,
        "A proper keyword-level share-of-voice comparison against the 31 named competitors "
        "(Dr. Reddy's, Aurobindo, Sun Pharma, Divi's, Laurus, Neuland, Hetero, Granules, Concord "
        "Biotech, and others) requires a SERP rank-tracking tool, which is not yet set up. Recommend "
        "adding this as a monthly line item once a tool (e.g. a rank tracker with competitor tracking) "
        "is in place — do not estimate competitor rankings without a real tool pull.",
    )


def build_work_completed(doc):
    heading(doc, "6. Work Completed This Period", size=14, before=8)
    section(doc, "aarisepharma.com", GRN)
    table(doc, ["Fix", "Scope", "Detail"], [
        [("Rewrote wrong-angle country posts", True, NAVY), ("5 posts", False, None),
         ("Brazil, Argentina, Chile, Colombia, Ecuador — importer-angle rewritten to supplier/export angle, FAQ schema added", False, None)],
        [("Resolved Peru/Paraguay cannibalization", False, None), ("2 posts", False, None),
         ("301 redirect into correct-angle pages via Code Snippets; old posts set to draft", False, None)],
        [("Rebuilt dispersible-tablets page", True, NAVY), ("1 page, 27K+ impressions/mo", False, None),
         ("Found and fixed broken WPBakery shortcode markup rendering as literal text to visitors; added FAQ schema", False, None)],
        [("Fixed content-corruption bug", False, None), ("16 posts", False, None),
         ('"Active [Pharmaceutical Ingredients]" truncation from a prior find-replace error', False, None)],
        [("llms.txt shipped", True, GRN), ("sitewide", False, None),
         ("Live at aarisepharma.com/llms.txt, content corrected to remove dead paths/over-claims", False, None)],
        [("Sitemap cleanup", False, None), ("14→8 child sitemaps", False, None),
         ("Removed 106 product-tag URLs plus slider/testimonial/brand/author bloat", False, None)],
        [("New country hub pages", True, UP), ("10 countries", False, None),
         ("Guatemala, Spain, Germany, Poland, Turkey, South Korea, Vietnam, Canada, Russia, Netherlands", False, None)],
        [("New tier-2/3 depth pages", True, UP), ("40 posts, 5 countries", False, None),
         ("Buy + Catalog pages (Steroid/Peptide/Pharma/Hormone) for Ecuador, Colombia, Chile, Brazil, Argentina", False, None)],
    ], [2.0, 1.3, 3.2])

    section(doc, "aarisehealthcare.com", BLU)
    table(doc, ["Fix", "Scope", "Detail"], [
        [("Fixed doubled-word bug", True, NAVY), ("24 posts", False, None),
         ('"Research Research Steroid/Peptide/Hormone Compound" typo', False, None)],
        [("RankMath meta backfill", True, NAVY), ("48 posts", False, None),
         ("Focus keyword, SEO title, meta description set on all existing country posts (previously unset)", False, None)],
        [("llms.txt content corrected", False, None), ("sitewide", False, None),
         ("Removed references to non-existent paths, stopped over-claiming country coverage", False, None)],
        [("New country hub pages", True, UP), ("15 countries", False, None),
         ("Brazil, Colombia, Chile, Argentina, Guatemala, Ecuador, Spain, Germany, Poland, Turkey, South Korea, Vietnam, Canada, Russia, Netherlands — RUO-hardened compliance language", False, None)],
    ], [2.0, 1.3, 3.2])


def build_promise_tracker(doc):
    heading(doc, "7. Promise Tracker", size=14, before=8)
    p(doc, "Every recommendation carried forward from prior reports, with an honest status — nothing "
           "gets silently re-recommended.", size=9.5, color=GREY)
    table(doc, ["Recommendation (from prior report)", "Status"], [
        [("Improve meta title on dispersible-tablets page for CTR", False, None),
         ("Done this period — page rebuilt with real content + FAQ schema", True, UP)],
        [("Install Code Snippets plugin on aarisepharma.com for llms.txt", False, None),
         ("Done — installed, llms.txt live", True, UP)],
        [("Add FAQ schema to dispersible-tablets post", False, None),
         ("Done this period", True, UP)],
        [("Create dedicated Third Party Manufacturing landing page", False, None),
         ("Not started — carrying forward", True, AMBER)],
        [("Healthcare: request re-indexing for country pages via GSC", False, None),
         ("Blocked on GSC access — carrying forward", True, DN)],
        [("Healthcare: remove footer Recent Products widget", False, None),
         ("Blocked on Sabhya/WP-admin — carrying forward, see Waiting on You", True, DN)],
        [("Healthcare: turn off pingbacks, fix author display name", False, None),
         ("Verify still needed — pending comments count showed 0 on last check", True, AMBER)],
    ], [3.8, 3.4])


def build_next_month_plan(doc):
    heading(doc, "8. Next Period Plan", size=14, before=8)
    numbered(doc, "Build tier-2/3 (Buy/Catalog) depth pages for the remaining 10 pharma countries "
                  "(Guatemala, Spain, Germany, Poland, Turkey, South Korea, Vietnam, Canada, Russia, Netherlands).",
             "1. Pharma depth, batch 2: ")
    numbered(doc, "Reassess GSC data on the 15 new Healthcare country hub pages; add tier-2/3 depth "
                  "only to the countries showing real movement.", "2. Healthcare depth (data-gated): ")
    numbered(doc, "Wire up a live GA4 Data API pull so Leads & Enquiries populates automatically next report.",
             "3. Close the leads-data gap: ")
    numbered(doc, "Set up a SERP rank-tracking tool for the competitor snapshot section.",
             "4. Close the competitor-data gap: ")
    numbered(doc, "Fix the sitewide Woodmart mega-menu content showing broken shortcode text — needs wp-admin access.",
             "5. Mega-menu bug: ")


def build_waiting_on_you(doc):
    heading(doc, "9. Waiting on You", size=14, before=8)
    p(doc, "Client/Sabhya-side actions — kept separate from the agency's own work above.", size=9.5, color=GREY)
    bullet(doc, "Create the Pharmacompass listing for aarisepharma.com — free, fastest available lead channel, not yet done.")
    bullet(doc, "Remove the footer Recent Products widget on aarisehealthcare.com (WP Admin → Appearance → Widgets).")
    bullet(doc, "Fix the sitewide broken mega-menu content on aarisepharma.com (WP Admin → Appearance → Menus) — "
                "or grant a real wp-admin login so we can fix it directly.")
    bullet(doc, "Confirm the GSC/GA4 service account (search-console@level-district-353301.iam.gserviceaccount.com) "
                "still has active access on both properties — needed to automate future reports.")


# ══════════════════════════════════════════════════════════════════════
def build_report():
    doc = new_doc()

    p(doc, "AANYA LABS · SEO REPORT", size=8, bold=True, color=GREY, before=8, after=4)
    heading(doc, "Aarise SEO — Combined Monthly Report", size=22, before=4, after=2)
    heading(doc, f"{CURRENT_LABEL}", size=14, before=0, after=4)
    p(doc, "aarisepharma.com + aarisehealthcare.com", size=11, color=GREY, before=0, after=16)
    rule(doc)
    for label, val in [
        ("Prepared by", "Aanya Labs"),
        ("Period", CURRENT_LABEL),
        ("Compared to", PREVIOUS_LABEL),
        ("Client", "LynqLeads / Rishikesh Sawant"),
        ("Contact", "Sabhya"),
    ]:
        mixed(doc, [(f"{label}:  ", True, GREY), (val, False, DARK)], size=10, before=3, after=3)
    doc.add_page_break()

    build_leads_section(doc)
    doc.add_page_break()

    pharma_gsc = load_gsc(PHARMA_MOM)
    build_search_performance(doc, "aarisepharma.com", pharma_gsc)
    doc.add_page_break()

    health_gsc = load_gsc(HEALTH_MOM)
    build_search_performance(doc, "aarisehealthcare.com", health_gsc)
    doc.add_page_break()

    pharma_status = [
        ("USA", "Live — full depth (tier 1-3)"), ("Mexico", "Live — full depth (tier 1-3)"),
        ("Peru", "Live — full depth (tier 1-3)"),
        ("Brazil", "Live — full depth (tier 1-3, published this period)"),
        ("Colombia", "Live — full depth (tier 1-3, published this period)"),
        ("Chile", "Live — full depth (tier 1-3, published this period)"),
        ("Argentina", "Live — full depth (tier 1-3, published this period)"),
        ("Guatemala", "Live — hub only (tier 1), depth queued next"),
        ("Ecuador", "Live — full depth (tier 1-3, published this period)"),
        ("Spain", "Live — hub only (tier 1), depth queued next"),
        ("Germany", "Live — hub only (tier 1), depth queued next"),
        ("Poland", "Live — hub only (tier 1), depth queued next"),
        ("Turkey", "Live — hub only (tier 1), depth queued next"),
        ("South Korea", "Live — hub only (tier 1), depth queued next"),
        ("Vietnam", "Live — hub only (tier 1), depth queued next"),
        ("Canada", "Live — hub only (tier 1), depth queued next"),
        ("Russia", "Live — hub only (tier 1), depth queued next"),
        ("Netherlands", "Live — hub only (tier 1), depth queued next"),
    ]
    health_status = [(c, "Live — hub (tier 1), RUO-compliant, published this period") for c in [
        "Brazil", "Colombia", "Chile", "Argentina", "Guatemala", "Ecuador", "Spain", "Germany",
        "Poland", "Turkey", "South Korea", "Vietnam", "Canada", "Russia", "Netherlands",
    ]] + [
        ("Mexico", "Live — full depth (tier 1-4)"), ("Peru", "Live — full depth (tier 1-4)"),
        ("USA", "Live — full depth (tier 1-4)"),
    ]
    build_country_tracker(doc, pharma_status, health_status)
    doc.add_page_break()

    pharma_ga4_summary = load_ga4_summary(PHARMA_GA4)
    pharma_ga4_pages = load_ga4_top_pages(PHARMA_GA4)
    pharma_ga4_sessions = load_ga4_session_source(PHARMA_GA4)
    pharma_ga4_cities = load_ga4_cities(PHARMA_GA4)
    build_ga4_detail(doc, "aarisepharma.com", pharma_ga4_summary, pharma_ga4_pages, pharma_ga4_sessions, pharma_ga4_cities)
    doc.add_page_break()

    health_ga4_summary = load_ga4_summary(HEALTH_GA4)
    health_ga4_pages = load_ga4_top_pages(HEALTH_GA4)
    health_ga4_sessions = load_ga4_session_source(HEALTH_GA4)
    health_ga4_cities = load_ga4_cities(HEALTH_GA4)
    build_ga4_detail(doc, "aarisehealthcare.com", health_ga4_summary, health_ga4_pages, health_ga4_sessions, health_ga4_cities)
    doc.add_page_break()

    build_competitor_snapshot(doc)
    doc.add_page_break()

    build_work_completed(doc)
    doc.add_page_break()

    build_promise_tracker(doc)
    doc.add_page_break()

    build_next_month_plan(doc)
    doc.add_page_break()

    build_waiting_on_you(doc)

    return doc


if __name__ == "__main__":
    doc = build_report()
    out_path = os.path.join(REPO, "reports", "monthly", "Aarise_Combined_SEO_Report_NEW_TEMPLATE.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
