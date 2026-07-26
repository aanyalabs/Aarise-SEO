from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
s = doc.sections[0]
s.left_margin = s.right_margin = Inches(1.0)
s.top_margin = s.bottom_margin = Inches(1.0)

NAVY  = RGBColor(0x0D,0x1B,0x2A); NAVY_HEX  = '0D1B2A'
GRN   = RGBColor(0x0A,0x7A,0x70); GRN_HEX   = '0A7A70'
BLU   = RGBColor(0x1A,0x5F,0xA8); BLU_HEX   = '1A5FA8'
GREY  = RGBColor(0x6B,0x7A,0x8D); GREY_HEX  = '6B7A8D'
DARK  = RGBColor(0x3D,0x4A,0x5C)
UP    = RGBColor(0x1A,0x7A,0x4A)
DN    = RGBColor(0xB8,0x30,0x30)

COLOR_HEX = {id(NAVY): NAVY_HEX, id(GRN): GRN_HEX, id(BLU): BLU_HEX, id(GREY): GREY_HEX}

def rgb_hex(color):
    return COLOR_HEX.get(id(color), GREY_HEX)

def p(text='', size=10, bold=False, color=None, before=2, after=4, align=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    if align: para.alignment = align
    r = para.add_run(text)
    r.font.size  = Pt(size)
    r.font.bold  = bold
    r.font.color.rgb = color or DARK
    return para

def mixed(parts, size=10, before=2, after=6):
    """parts = list of (text, bold, color)"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    for text, bold, color in parts:
        r = para.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.color.rgb = color or DARK
    return para

def heading(text, size=18, color=None, before=14, after=4):
    return p(text, size=size, bold=True, color=color or NAVY, before=before, after=after)

def rule(color='CCCCCC'):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6'); bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)
    return para

def section(text, color=None):
    rule(rgb_hex(color or GREY))
    return p(text.upper(), size=8, bold=True, color=color or GREY, before=2, after=6)

def shd(cell, hex6):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    e = OxmlElement('w:shd')
    e.set(qn('w:val'),'clear'); e.set(qn('w:color'),'auto'); e.set(qn('w:fill'),hex6)
    tcPr.append(e)

def table(headers, rows, widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    # header
    for i,h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        shd(c,'E8ECF2')
        para = c.paragraphs[0]
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(2)
        r = para.add_run(h)
        r.font.size = Pt(8.5); r.font.bold = True; r.font.color.rgb = GREY
        if i > 0: para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # data
    for ri, row in enumerate(rows):
        bg = 'FFFFFF' if ri%2==0 else 'F7F8FA'
        for ci, cell in enumerate(row):
            if isinstance(cell, tuple):
                txt = str(cell[0])
                bold  = cell[1] if len(cell)>1 else False
                color = cell[2] if len(cell)>2 else None
            else:
                txt, bold, color = str(cell), False, None
            c = tbl.rows[ri+1].cells[ci]
            shd(c, bg)
            para = c.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            r = para.add_run(txt)
            r.font.size = Pt(9); r.font.bold = bool(bold)
            r.font.color.rgb = color if isinstance(color, RGBColor) else DARK
            if ci > 0: para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if widths:
        for i,w in enumerate(widths):
            for row in tbl.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl

def bullet(text, lead=''):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    if lead:
        r = para.add_run(lead); r.font.bold=True; r.font.size=Pt(10); r.font.color.rgb=NAVY
    r2 = para.add_run(text); r2.font.size=Pt(10); r2.font.color.rgb=DARK

def numbered(text, lead=''):
    para = doc.add_paragraph(style='List Number')
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(2)
    if lead:
        r = para.add_run(lead); r.font.bold=True; r.font.size=Pt(10); r.font.color.rgb=NAVY
    r2 = para.add_run(text); r2.font.size=Pt(10); r2.font.color.rgb=DARK

def callout(bold_text, body, color=GRN):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(8)
    para.paragraph_format.left_indent  = Inches(0.15)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    lft  = OxmlElement('w:left')
    lft.set(qn('w:val'),'single'); lft.set(qn('w:sz'),'18')
    lft.set(qn('w:color'), rgb_hex(color))
    pBdr.append(lft); pPr.append(pBdr)
    r1 = para.add_run(bold_text+' '); r1.font.bold=True; r1.font.size=Pt(9.5); r1.font.color.rgb=color
    r2 = para.add_run(body); r2.font.size=Pt(9.5); r2.font.color.rgb=DARK

# ══════════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════════
p('AANYA LABS · SEO REPORT', size=8, bold=True, color=GREY, before=8, after=4)
heading('Monthly SEO Performance Report', size=24, before=4, after=2)
heading('June 2026', size=18, before=0, after=4)
p('aarisehealthcare.com & aarisepharma.com', size=11, color=GREY, before=0, after=16)
rule()
for label, val in [
    ('Prepared by','Priyansh Arora, Aanya Labs'),
    ('Report Date','July 23, 2026'),
    ('Period','June 20 – July 22, 2026'),
    ('Compared to','May 20 – June 19, 2026'),
    ('Client','LynqLeads / Rishikesh Sawant'),
    ('Contact','Sabhya'),
]:
    mixed([(f'{label}:  ', True, GREY),(val, False, DARK)], size=10, before=3, after=3)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════
heading('Executive Summary', size=14, before=8)
section('Overview')

mixed([
    ('Aarise Pharma', True, NAVY),
    (' had a strong month. Impressions grew by ', False, None),
    ('+9,756 (+20%) to 59,628', True, NAVY),
    (' and clicks rose ', False, None),
    ('+44 to 439', True, NAVY),
    ('. 800 of 1,340 tracked keywords now rank on Page 1 (60%). The dispersible tablets cluster continues to dominate organic rankings, and GA4 shows 986 organic sessions alongside ', False, None),
    ('16 sessions attributed to "AI Assistant"', True, GRN),
    (' — an early signal that AI search tools are beginning to surface Aarise Pharma in results.', False, None),
], size=10)

mixed([
    ('Aarise Healthcare', True, NAVY),
    (' is in an early-growth phase. Impressions grew ', False, None),
    ('+81 (+14%) to 655', True, NAVY),
    (' and average position improved dramatically from ', False, None),
    ('62.7 to 45.4', True, NAVY),
    (' — a 17-position jump MoM. GA4 reports 232 active users with an exceptional ', False, None),
    ('138-second average engagement time', True, NAVY),
    ('. Clicks will follow once Google re-crawls the updated pages and positions break into Page 2.', False, None),
], size=10)

mixed([
    ('Technical work completed this month: ', True, NAVY),
    ('286 duplicate pages noindexed, 721 meta descriptions added, FAQ schema on 42 supplier posts and 4 country pages, WooCommerce pages noindexed, and ', False, None),
    ('llms.txt published at aarisehealthcare.com/llms.txt', True, GRN),
    (' for AI search visibility.', False, None),
], size=10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# AARISE PHARMA
# ══════════════════════════════════════════════════════════════════════
heading('Aarise Pharma', size=18, color=GRN, before=8)
p('aarisepharma.com', size=10, color=GREY, before=0, after=2)
rule('0A7A70')

section('Search Console — This Month\'s Performance', GRN)
table(
    ['Metric','This Month','vs Prev Period'],
    [
        [('Clicks',True,NAVY),('439',True,NAVY),('+44 ▲',True,UP)],
        [('Impressions',True,NAVY),('59,628',True,NAVY),('+9,756 (+20%) ▲',True,UP)],
        [('CTR',False,None),('0.74%',False,None),('0.79% prev ▼',False,DN)],
        [('Avg. Position',False,None),('6.1',False,None),('5.3 prev (more KWs indexed)',False,GREY)],
        [('Keywords Tracked',False,None),('1,340',False,None),('—',False,GREY)],
        [('Page 1 Keywords',True,NAVY),('800',True,UP),('60% of total ▲',True,UP)],
    ],
    [2.8,1.4,2.3]
)
p('Impressions grew 20% MoM, reflecting continued index expansion. The slight position drop from 5.3 → 6.1 is expected as more long-tail keywords enter the index — it does not represent a ranking decline on core terms.', size=10)

section('Top Keywords (Search Console)', GRN)
p('The following keywords drove the most impressions and clicks this month:', size=10, after=4)
table(
    ['Keyword','Clicks','Impressions','CTR','Position','vs Prev Pos'],
    [
        [('dispersible tablets',True,NAVY),('43',False,None),('10,762',False,None),('0.40%',False,None),('4.0',True,NAVY),('−0.2 ▲',False,UP)],
        [('dispersible tablet',False,None),('11',False,None),('3,915',False,None),('0.28%',False,None),('5.0',False,None),('−0.4 ▲',False,UP)],
        [('dispersible',False,None),('7',False,None),('2,929',False,None),('0.24%',False,None),('6.8',False,None),('+0.8 ▼',False,DN)],
        [('dispersible tablets uses',False,None),('4',False,None),('1,836',False,None),('0.22%',False,None),('2.6',False,None),('+0.7 ▼',False,DN)],
        [('dispersible tablets meaning',False,None),('9',False,None),('1,390',False,None),('0.65%',False,None),('2.7',False,None),('+0.6 ▼',False,DN)],
        [('dextro',False,None),('5',False,None),('1,209',False,None),('0.41%',False,None),('9.4',False,None),('0.0',False,GREY)],
        [('dispersible tablet uses',False,None),('3',False,None),('1,133',False,None),('0.26%',False,None),('3.0',False,None),('0.0',False,GREY)],
        [('medicated toothpaste',False,None),('4',False,None),('743',False,None),('0.54%',False,None),('1.8',False,None),('0.0',False,GREY)],
        [('dispersible tablets benefits',False,None),('3',False,None),('520',False,None),('0.58%',False,None),('3.3',False,None),('+0.2',False,GREY)],
        [('aarise pharmaceuticals',True,NAVY),('119',False,None),('235',False,None),('50.6%',True,UP),('1.0',True,UP),('−0.4 ▲',False,UP)],
        [('top pharma API supplier from India',False,None),('23',False,None),('77',False,None),('29.9%',False,None),('1.4',True,UP),('stable',False,GREY)],
        [('dispersible tablets advantages',False,None),('15',False,None),('74',False,None),('20.3%',False,None),('3.2',False,None),('−1.9 ▲',False,UP)],
        [('aarise pharma',True,NAVY),('15',False,None),('74',False,None),('20.3%',True,UP),('1.1',True,UP),('−0.3 ▲',False,UP)],
        [('third party manufacturing',False,None),('~5',False,None),('~2,900',False,None),('0.17%',False,None),('10.4',False,None),('−0.3',False,GREY)],
    ],
    [2.2,0.6,0.9,0.65,0.75,0.8]
)
mixed([
    ('Branded keywords', True, NAVY),
    (' (aarise pharmaceuticals, aarise pharma) continue performing at extremely high CTR — 50.6% and 20.3%. The dispersible tablets cluster remains the dominant organic driver. ', False, None),
    ('"Top pharma API supplier from India"', True, NAVY),
    (' at position 1.4 is a highly commercial keyword now performing exceptionally well.', False, None),
], size=10)

section('Top Pages (Search Console)', GRN)
table(
    ['Page','Clicks','Impressions','CTR','Position'],
    [
        [('What is Dispersible Tablets?',True,NAVY),('~30',False,None),('~2,200',False,None),('1.4%',False,None),('5.6',False,None)],
        [('Homepage',False,None),('~30',False,None),('~300',False,None),('10.1%',False,None),('2.1',False,None)],
        [('Products Page',False,None),('~14',False,None),('~3,100',False,None),('0.45%',False,None),('7.8',False,None)],
        [('Contact Us',False,None),('~11',False,None),('~296',False,None),('3.7%',False,None),('3.4',False,None)],
        [('Third Party Manufacturing',False,None),('~5',False,None),('~2,900',False,None),('0.17%',False,None),('10.4',False,None)],
    ],
    [2.8,0.7,0.9,0.7,0.8]
)
p('The "What is Dispersible Tablets" blog post continues to be the strongest organic page. Third Party Manufacturing holds 2,900 impressions at position 10.4 — moving this to Page 1 is a near-term priority.', size=10)

section('Google Analytics — Overview', GRN)
table(
    ['Metric','Value'],
    [
        [('New Users',False,None),('~4,858',False,None)],
        [('Organic Search Sessions',False,None),('986',False,None)],
        [('Direct Sessions',False,None),('4,094',False,None)],
        [('Referral Sessions',False,None),('45',False,None)],
        [('Organic Social',False,None),('30',False,None)],
        [('AI Assistant Sessions',True,GRN),('16 ★',True,GRN)],
        [('Total Page Views',False,None),('6,651',False,None)],
    ],
    [3.5,3.0]
)
callout('AI Traffic Milestone:', 'GA4 is now reporting 16 sessions under the "AI Assistant" channel — ChatGPT, Gemini, or Perplexity are surfacing Aarise Pharma in their responses and users are clicking through. This is the first month this signal appears and will grow as our llms.txt and entity-building work takes effect.', GRN)

section('Traffic Sources (GA4)', GRN)
table(
    ['Source / Channel','Sessions'],
    [
        [('Direct / (none)',False,None),('4,094',False,None)],
        [('Google / Organic',False,None),('986',False,None)],
        [('Referral',False,None),('45',False,None)],
        [('Organic Social',False,None),('30',False,None)],
        [('Unassigned',False,None),('26',False,None)],
        [('AI Assistant',True,GRN),('16',True,GRN)],
    ],
    [3.5,3.0]
)

section('Top Pages (GA4)', GRN)
table(
    ['Page','Views'],
    [
        [('Home',False,None),('977',False,None)],
        [('Products',False,None),('606',False,None)],
        [('Contact Us',False,None),('239',False,None)],
        [('What is Dispersible Tablets?',False,None),('182',False,None)],
        [('Third Party Manufacturing',False,None),('113',False,None)],
    ],
    [3.5,3.0]
)

section('Ranking Summary', GRN)
table(
    ['Ranking','Count','% of Total'],
    [
        [('Page 1 (Pos 1–10)',True,NAVY),('800',True,UP),('60%',True,UP)],
        [('Page 2 (Pos 11–20)',False,None),('122',False,None),('9%',False,None)],
        [('Page 3 (Pos 21–30)',False,None),('54',False,None),('4%',False,None)],
        [('Page 4+',False,None),('174',False,None),('13%',False,None)],
        [('Not Yet Indexed',False,None),('190',False,DN),('14%',False,DN)],
        [('TOTAL',True,NAVY),('1,340',True,NAVY),('100%',True,NAVY)],
    ],
    [3.5,1.0,2.0]
)
p('60% of tracked keywords ranking on Page 1 is an excellent result. The 190 not-yet-indexed keywords represent a growth opportunity — refreshed sitemap submission will bring these into rankings.', size=10)

section('Summary & Recommendations', GRN)
p('Key Takeaways:', size=10, bold=True, color=NAVY, after=2)
bullet('439 clicks and 59,628 impressions — +20% impression growth MoM.', '▲ ')
bullet('800 keywords on Page 1 — 60% of all tracked keywords.', '✓ ')
bullet('16 AI Assistant sessions — the site is being cited by AI tools for the first time.', '★ ')
bullet('Dispersible tablets cluster — 10,762 impressions on primary keyword alone.', '▲ ')
bullet('Branded keywords (aarise pharmaceuticals, aarise pharma) at 50%+ CTR.', '✓ ')
bullet('Third Party Manufacturing at position 10.4 — one push away from Page 1.', '→ ')
p('Recommendations:', size=10, bold=True, color=NAVY, before=8, after=2)
numbered('Improve meta title on dispersible tablets page to push CTR above 1% — currently 0.40% on 10,762 impressions = ~100 uncaptured clicks/month.', 'Improve CTR: ')
numbered('Create dedicated Third Party Manufacturing landing page — 2,900 impressions at position 10.4.', 'New landing page: ')
numbered('Install Code Snippets plugin on aarisepharma.com to activate /llms.txt — page is already published.', 'Activate llms.txt: ')
numbered('Add FAQ schema to the dispersible tablets blog post for Google AI Overview eligibility.', 'FAQ schema: ')
numbered('Monitor AI Assistant sessions monthly — new channel to track and grow.', 'Track AI channel: ')
numbered('Submit refreshed sitemap to bring 190 unindexed keywords into rankings.', 'Sitemap: ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# AARISE HEALTHCARE
# ══════════════════════════════════════════════════════════════════════
heading('Aarise Healthcare', size=18, color=BLU, before=8)
p('aarisehealthcare.com', size=10, color=GREY, before=0, after=2)
rule('1A5FA8')

section('Search Console — This Month\'s Performance', BLU)
table(
    ['Metric','This Month','vs Prev Period'],
    [
        [('Clicks',False,None),('0',False,None),('Positions still improving',False,GREY)],
        [('Impressions',True,NAVY),('655',True,NAVY),('+81 (+14%) ▲',True,UP)],
        [('CTR',False,None),('0.0%',False,None),('Requires Page 1–2 positions',False,GREY)],
        [('Avg. Position',True,NAVY),('45.4',True,UP),('62.7 prev — improved 17 pos ▲',True,UP)],
        [('Keywords Tracked',False,None),('269',False,None),('—',False,GREY)],
        [('Page 1+2 Keywords',True,NAVY),('15',True,UP),('6 Page 1 + 9 Page 2 ▲',True,UP)],
    ],
    [2.8,1.4,2.3]
)
mixed([
    ('Average position improved 17 positions (62.7 to 45.4) — the largest MoM gain since launch. This is a direct result of the technical fixes completed this month. Google has not fully re-crawled the updated pages yet — the next 30 days should show this convert into Page 1 rankings and first clicks.', False, None),
], size=10)

section('Top Keywords (Search Console)', BLU)
p('The following keywords show the strongest position signals this month:', size=10, after=4)
table(
    ['Keyword','Clicks','Impressions','CTR','Position','vs Prev Pos'],
    [
        [('research compounds',True,NAVY),('0',False,None),('179',False,None),('0%',False,None),('15.3',False,None),('▲ +25.5',True,UP)],
        [('compound research',False,None),('0',False,None),('31',False,None),('0%',False,None),('32.4',False,None),('▲ +15.8',True,UP)],
        [('compounds for research',False,None),('0',False,None),('24',False,None),('0%',False,None),('12.6',False,None),('▲ +4.4',True,UP)],
        [('laboratory research compounds',False,None),('0',False,None),('36',False,None),('0%',False,None),('50.0',False,None),('▲ +10.6',True,UP)],
        [('research compound',True,NAVY),('0',False,None),('5',False,None),('0%',False,None),('7.8 ✓',True,UP),('▲ +16.7',True,UP)],
        [('aarise',True,NAVY),('0',False,None),('5',False,None),('0%',False,None),('9.6 ✓',True,UP),('New',True,UP)],
        [('coa report',False,None),('0',False,None),('3',False,None),('0%',False,None),('7.0 ✓',True,UP),('▲ +14.5',True,UP)],
        [('arise healthcare',False,None),('0',False,None),('8',False,None),('0%',False,None),('18.4',False,None),('▲ +21.5',True,UP)],
        [('compounds for research',False,None),('0',False,None),('24',False,None),('0%',False,None),('12.6',False,None),('▲ +4.4',True,UP)],
        [('semaglutide peptide supplier',False,None),('0',False,None),('4',False,None),('0%',False,None),('42.5',False,None),('New',True,UP)],
        [('research peptides',False,None),('0',False,None),('7',False,None),('0%',False,None),('69.1',False,None),('New',True,UP)],
    ],
    [2.2,0.6,0.9,0.65,0.85,0.8]
)
mixed([
    ('"Research compounds" at position 15.3 (up from 40.8) is the most important keyword — just outside Page 1. ', False, None),
    ('"Research compound" at 7.8 and "aarise" at 9.6 are both on Page 1.', True, NAVY),
    (' Several new keywords entered the index for the first time this month.', False, None),
], size=10)
callout('Key Signal:', '"Research compound" is now at position 7.8 — Page 1. "Compounds for research" is at 12.6, just below Page 1. One targeted piece of content on these terms could push both into top 10 and deliver the site\'s first consistent organic clicks.', BLU)

section('Google Analytics — Overview', BLU)
table(
    ['Metric','Value'],
    [
        [('Active Users',False,None),('232',False,None)],
        [('New Users',False,None),('216',False,None)],
        [('Avg. Engagement Time',True,NAVY),('138 seconds',True,UP)],
        [('Total Events',False,None),('3,870',False,None)],
        [('Checkout Page Views',True,NAVY),('13  (5.6% checkout rate)',True,UP)],
    ],
    [3.5,3.0]
)
mixed([
    ('138-second average engagement time', True, NAVY),
    (' is exceptional — nearly 6x higher than aarisepharma.com. Users who find the site are deeply reading product pages, indicating high buyer intent. The issue is purely discovery (search positions), not content quality or conversion. 13 checkout page views from 232 users = ', False, None),
    ('5.6% checkout rate', True, NAVY),
    (' — high for a research compound site at this stage.', False, None),
], size=10)

section('Top Pages (GA4)', BLU)
table(
    ['Page','Views','Active Users','Bounce Rate'],
    [
        [('Peptides',True,NAVY),('253',False,None),('86',False,None),('21%',False,None)],
        [('Steroids',False,None),('242',False,None),('74',False,None),('25%',False,None)],
        [('Home',False,None),('172',False,None),('88',False,None),('31%',False,None)],
        [('Shop',False,None),('158',False,None),('27',False,None),('20%',False,None)],
        [('Hormones',False,None),('76',False,None),('43',False,None),('6%',False,None)],
        [('Pharma',False,None),('75',False,None),('40',False,None),('37%',False,None)],
        [('Tirzepatide',False,None),('17',False,None),('8',False,None),('29%',False,None)],
        [('Checkout',True,NAVY),('13',False,None),('11',False,None),('17%',False,None)],
    ],
    [2.3,1.0,1.3,1.8]
)
p('Peptides and Steroids are the top two category pages. The Hormones page has an exceptionally low 6% bounce rate. Checkout views (13) confirm purchase intent is present.', size=10)

section('Technical Work Completed This Month', BLU)
table(
    ['Fix','Pages','Impact'],
    [
        [('Duplicate pages noindexed ("-2" suffix)',True,NAVY),('286',False,None),('Consolidates ranking signals to canonical pages',False,None)],
        [('Meta descriptions added — posts',False,None),('91',False,None),('Improves CTR and search snippet quality',False,None)],
        [('Meta descriptions added — product pages',False,None),('293',False,None),('Unique keyword-targeted copy on all products',False,None)],
        [('FAQ schema added (supplier posts)',False,None),('42',False,None),('Google AI Overview eligibility',False,None)],
        [('Country pages fully rewritten',True,NAVY),('4',False,None),('Paraguay, Peru, Mexico, US — 200→800+ words + FAQ schema',False,None)],
        [('WooCommerce system pages noindexed',False,None),('5',False,None),('Crawl budget reclaimed',False,None)],
        [('llms.txt published',True,GRN),('1',False,None),('Live at aarisehealthcare.com/llms.txt — AI search visibility',False,GRN)],
    ],
    [2.3,0.6,3.5]
)

section('Ranking Summary', BLU)
table(
    ['Ranking','Count','% of Total'],
    [
        [('Page 1 (Pos 1–10)',True,NAVY),('6',True,UP),('2%',True,UP)],
        [('Page 2 (Pos 11–20)',False,None),('9',False,None),('3%',False,None)],
        [('Page 3 (Pos 21–30)',False,None),('6',False,None),('2%',False,None)],
        [('Page 4+',False,None),('142',False,None),('53%',False,None)],
        [('Not Yet Indexed',False,None),('106',False,DN),('39%',False,DN)],
        [('TOTAL',True,NAVY),('269',True,NAVY),('100%',True,NAVY)],
    ],
    [3.5,1.0,2.0]
)
p('The high not-yet-indexed percentage (39%) reflects the site\'s early stage. This month\'s technical fixes directly address this — requesting re-indexing via GSC URL Inspection is the most immediate action.', size=10)

section('Summary & Recommendations', BLU)
p('Key Takeaways:', size=10, bold=True, color=NAVY, after=2)
bullet('Average position improved 17 positions (62.7 → 45.4) — largest MoM gain to date.', '▲ ')
bullet('"Research compound" is now at position 7.8 — Page 1. "Aarise" at 9.6 — Page 1.', '✓ ')
bullet('655 impressions, +14% growth, driven by technical improvements.', '▲ ')
bullet('232 active users with 138-second average engagement — exceptional content quality.', '✓ ')
bullet('721 pages now have unique meta descriptions — foundational fix compounding over time.', '✓ ')
bullet('286 duplicate pages noindexed — ranking signals now consolidate on canonical pages.', '✓ ')
bullet('llms.txt is live at aarisehealthcare.com/llms.txt — AI crawlers can discover the site.', '★ ')
p('Recommendations:', size=10, bold=True, color=NAVY, before=8, after=2)
numbered('In GSC URL Inspection, request indexing for the Paraguay, Peru, Mexico, and US country pages. Submit sitemap: aarisehealthcare.com/sitemap_index.xml. Highest-leverage action right now.', 'Request re-indexing immediately: ')
numbered('"Research compounds" is at position 15.3 with 179 impressions. A guide targeting this term with FAQ schema could push it to Page 1 and deliver first consistent clicks.', 'Publish article on "research compounds": ')
numbered('"Compounds for research" at position 12.6 — just off Page 1. Internal link cluster from existing pages should push it over.', 'Target "compounds for research": ')
numbered('Add Product schema to individual product pages for rich snippet eligibility.', 'Product schema: ')
numbered('Continue internal links from blog posts to high-value product pages (Tirzepatide, Semaglutide, BPC-157).', 'Internal linking: ')
numbered('Remove footer Recent Products widget, turn off pingbacks, fix author display name, trash 52 pending spam comments.', 'Sabhya admin actions: ')

# ── Save ──
out = r"C:\Users\priya\Downloads\Aarise_SEO_Monthly_Report_June_2026.docx"
doc.save(out)
print(f"Done: {out}")
